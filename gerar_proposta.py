"""
Gerador de Propostas Comerciais em DOCX
Azevedo Contabilidade
"""

import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ===== CORES =====
GOLD = RGBColor(0xB8, 0x96, 0x0C)
DARK = RGBColor(0x1A, 0x27, 0x44)
GRAY = RGBColor(0x66, 0x66, 0x66)
LGRAY = RGBColor(0x99, 0x99, 0x99)
TEXT_COLOR = RGBColor(0x4A, 0x4A, 0x4A)
LIGHT_TEXT = RGBColor(0x5A, 0x5A, 0x5A)


def set_cell_shading(cell, color_hex):
    """Define cor de fundo da celula"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=55, bottom=55, left=100, right=100):
    """Define margens internas da celula"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_paragraph(doc, text="", font_size=10.5, color=None, bold=False, alignment=None, space_after=None, space_before=None, font_name="Arial"):
    """Adiciona paragrafo formatado"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(font_size * 1.5)

    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.font.color.rgb = color or TEXT_COLOR
    run.font.bold = bold
    return p


def add_horizontal_line(doc, color="B8960C", size=4):
    """Adiciona linha horizontal usando borda inferior do paragrafo"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    # Add border via paragraph properties - must use the border property directly
    p_fmt = p.paragraph_format
    # Use the low-level approach with proper element ordering
    pPr = p._p.get_or_add_pPr()
    # Remove any existing pBdr
    for existing in pPr.findall(qn('w:pBdr')):
        pPr.remove(existing)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    # Insert pBdr after spacing elements but before jc
    # Find proper insertion point
    inserted = False
    for child in list(pPr):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('jc', 'rPr', 'sectPr'):
            child.addprevious(pBdr)
            inserted = True
            break
    if not inserted:
        pPr.append(pBdr)


def fc(valor):
    """Formata valor como R$ 1.234,56"""
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def gerar_docx(dados):
    """
    Gera proposta comercial em DOCX.

    dados = {
        tratamento, nome, telefone, email, vendedor, introducao,
        servicos: [{descricao, valor, periodicidade}],
        desconto_pct, pix_cnpj, pix_titular, observacao,
        incluir_doc, texto_doc, logo_path
    }
    """
    doc = Document()

    # ===== PAGE SETUP =====
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ===== DEFAULT FONT =====
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    style.font.color.rgb = TEXT_COLOR
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ===== LOGO + HEADER =====
    logo_path = dados.get("logo_path", "")

    # Header com logo e info da empresa lado a lado via tabela invisivel
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove bordas
    for row in header_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    # Logo cell
    logo_cell = header_table.cell(0, 0)
    logo_cell.width = Cm(4)
    logo_p = logo_cell.paragraphs[0]
    if os.path.exists(logo_path):
        logo_p.add_run().add_picture(logo_path, width=Cm(3))

    # Info cell
    info_cell = header_table.cell(0, 1)
    info_cell.width = Cm(12)

    # Clear default paragraph
    info_p = info_cell.paragraphs[0]
    info_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = info_p.add_run("AZEVEDO CONTABILIDADE")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = DARK
    run.font.name = "Arial"

    for line in [
        "Av. Antônio Basílio, 3025 | Lagoa Nova",
        "Natal – RN | CEP 59056-500",
        "CNPJ: 35.304.872/0001-28"
    ]:
        p = info_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(line)
        r.font.size = Pt(8)
        r.font.color.rgb = LGRAY
        r.font.name = "Arial"

    # ===== GOLD LINE =====
    add_horizontal_line(doc, "B8960C", 6)

    # ===== DATE =====
    meses = ["janeiro","fevereiro","março","abril","maio","junho","julho","agosto","setembro","outubro","novembro","dezembro"]
    hoje = datetime.now()
    data_str = f"NATAL/RN, {hoje.day} DE {meses[hoje.month-1].upper()} DE {hoje.year}"
    add_paragraph(doc, data_str, font_size=9, color=LGRAY, space_before=8, space_after=4)

    # ===== RECIPIENT =====
    add_paragraph(doc, f"{dados['tratamento']} {dados['nome']}", font_size=14, color=DARK, bold=True, space_after=8)

    # ===== INTRO =====
    intro = dados.get("introducao", "prestação de serviços contábeis")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(17)

    r1 = p.add_run("Apresentamos a ")
    r1.font.size = Pt(10)
    r1.font.color.rgb = TEXT_COLOR
    r1.font.name = "Arial"

    r2 = p.add_run("V. Sa.")
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.color.rgb = TEXT_COLOR
    r2.font.name = "Arial"

    r3 = p.add_run(" nossa proposta referente a ")
    r3.font.size = Pt(10)
    r3.font.color.rgb = TEXT_COLOR
    r3.font.name = "Arial"

    r4 = p.add_run(intro)
    r4.font.size = Pt(10)
    r4.font.bold = True
    r4.font.color.rgb = TEXT_COLOR
    r4.font.name = "Arial"

    r5 = p.add_run(", conforme condições a seguir:")
    r5.font.size = Pt(10)
    r5.font.color.rgb = TEXT_COLOR
    r5.font.name = "Arial"

    # ===== LINE =====
    add_horizontal_line(doc, "E8E8E8", 1)

    # ===== NOSSA EMPRESA =====
    add_paragraph(doc, "NOSSA EMPRESA", font_size=7.5, color=GOLD, bold=True, space_before=6, space_after=4)

    nossa_empresa = (
        "A Azevedo Contabilidade é um escritório especializado em contabilidade estratégica "
        "e planejamento tributário, com atuação no Rio Grande do Norte e atendimento a empresas "
        "em diferentes estados do país. Nosso trabalho é voltado a apoiar empresários na tomada "
        "de decisões, garantindo organização contábil, segurança fiscal e eficiência tributária. "
        "Mais do que cumprir obrigações legais, buscamos atuar como parceiros do crescimento dos "
        "nossos clientes, oferecendo informações confiáveis, tempestivas e relevantes para a gestão "
        "do negócio. Para isso, contamos com uma equipe qualificada, processos estruturados e o uso "
        "de tecnologia que assegura agilidade, precisão e transparência nos serviços prestados. "
        "Nossa missão é proporcionar uma contabilidade que gere valor para a empresa, contribuindo "
        "para redução de riscos, melhoria da gestão e otimização da carga tributária."
    )
    p = add_paragraph(doc, nossa_empresa, font_size=8.5, color=LIGHT_TEXT, space_after=4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(14)

    # ===== LINE =====
    add_horizontal_line(doc, "E8E8E8", 1)

    # ===== SECTION 1: SERVICOS =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("01  ")
    r1.font.size = Pt(10)
    r1.font.bold = True
    r1.font.color.rgb = GOLD
    r1.font.name = "Arial"
    r2 = p.add_run("SERVIÇOS CONTRATADOS")
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.color.rgb = DARK
    r2.font.name = "Arial"

    add_horizontal_line(doc, "E8E8E8", 1)

    for svc in dados["servicos"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.8)
        r = p.add_run("◆  ")
        r.font.size = Pt(7)
        r.font.color.rgb = GOLD
        r.font.name = "Arial"
        r2 = p.add_run(svc["descricao"])
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
        r2.font.name = "Arial"

    # ===== SECTION 2: TABLE =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("02  ")
    r1.font.size = Pt(10)
    r1.font.bold = True
    r1.font.color.rgb = GOLD
    r1.font.name = "Arial"
    r2 = p.add_run("HONORÁRIOS E INVESTIMENTO")
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.color.rgb = DARK
    r2.font.name = "Arial"

    add_horizontal_line(doc, "E8E8E8", 1)

    # Build table
    headers = ["Descrição", "Valor Unit.", "Desconto", "Periodicidade", "Total"]
    num_rows = len(dados["servicos"]) + 2  # header + data + total
    table = doc.add_table(rows=num_rows, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    for j, header_text in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, "F7F7F7")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header_text)
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = DARK
        r.font.name = "Arial"

    # Data rows
    pct = dados.get("desconto_pct", 0.10)
    total_geral = 0

    for i, svc in enumerate(dados["servicos"]):
        real = svc["valor"]
        if pct > 0 and real > 0:
            unit_inf = real / (1 - pct)
            desconto = unit_inf - real
        else:
            unit_inf = real
            desconto = 0
        total = real
        total_geral += total

        row_data = [
            svc["descricao"],
            fc(unit_inf),
            f"-{fc(desconto)}" if desconto > 0 else "-",
            svc["periodicidade"],
            fc(total)
        ]

        bg = "FAFAFA" if (i % 2 == 1) else "FFFFFF"
        for j, txt in enumerate(row_data):
            cell = table.cell(i + 1, j)
            set_cell_shading(cell, bg)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(txt)
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
            r.font.name = "Arial"

    # Total row
    total_row_idx = len(dados["servicos"]) + 1

    # Merge first 4 cells for "TOTAL"
    total_cell = table.cell(total_row_idx, 0)
    total_cell.merge(table.cell(total_row_idx, 3))
    set_cell_shading(total_cell, "FAF6E6")
    set_cell_margins(total_cell)
    p = total_cell.paragraphs[0]
    r = p.add_run("TOTAL")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = DARK
    r.font.name = "Arial"

    val_cell = table.cell(total_row_idx, 4)
    set_cell_shading(val_cell, "FAF6E6")
    set_cell_margins(val_cell)
    p = val_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(fc(total_geral))
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = DARK
    r.font.name = "Arial"

    # Set table borders using proper element ordering
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    # Remove existing borders if any
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:start w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:end w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="F0F0F0"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="F0F0F0"/>'
        f'</w:tblBorders>'
    )
    # Insert after tblStyle if present, otherwise as first child
    tblStyle = tblPr.find(qn('w:tblStyle'))
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblW.addnext(borders)
    elif tblStyle is not None:
        tblStyle.addnext(borders)
    else:
        tblPr.insert(0, borders)

    # ===== OBSERVACAO =====
    obs = dados.get("observacao", "")
    if obs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run("Observação: ")
        r1.font.size = Pt(8.5)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
        r1.font.name = "Arial"
        r2 = p.add_run(obs)
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = LIGHT_TEXT
        r2.font.name = "Arial"

    # ===== PIX =====
    add_paragraph(doc, "", space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Dados para pagamento via PIX:")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = DARK
    r.font.name = "Arial"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run("Chave CNPJ: ")
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
    r1.font.name = "Arial"
    r2 = p.add_run(dados["pix_cnpj"])
    r2.font.size = Pt(8.5)
    r2.font.bold = True
    r2.font.color.rgb = DARK
    r2.font.name = "Arial"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run("Titular: ")
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
    r1.font.name = "Arial"
    r2 = p.add_run(dados["pix_titular"])
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = DARK
    r2.font.name = "Arial"

    # ===== DOCUMENTACAO =====
    if dados.get("incluir_doc") and dados.get("texto_doc"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run("03  ")
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = GOLD
        r1.font.name = "Arial"
        r2 = p.add_run("DOCUMENTAÇÃO NECESSÁRIA")
        r2.font.size = Pt(10)
        r2.font.bold = True
        r2.font.color.rgb = DARK
        r2.font.name = "Arial"

        add_horizontal_line(doc, "E8E8E8", 1)

        p = add_paragraph(doc, dados["texto_doc"], font_size=9.5, color=TEXT_COLOR, space_after=8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = Pt(16)

    # ===== SIGNATURE (page 2) =====
    # Linha dourada separadora
    add_horizontal_line(doc, "B8960C", 6)

    # Espaço generoso antes da assinatura
    add_paragraph(doc, "", space_after=40)

    # Bloco de assinatura
    add_paragraph(doc, "Atenciosamente,", font_size=10, color=TEXT_COLOR,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # Linha para assinatura
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("_" * 50)
    r.font.size = Pt(10)
    r.font.color.rgb = LGRAY
    r.font.name = "Arial"

    add_paragraph(doc, "Azevedo Contabilidade", font_size=12, color=DARK, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "Contabilidade Estratégica & Planejamento Tributário", font_size=8.5, color=LGRAY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # Nome do vendedor se disponível
    vendedor = dados.get("vendedor", "")
    if vendedor:
        add_paragraph(doc, "", space_after=20)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("_" * 50)
        r.font.size = Pt(10)
        r.font.color.rgb = LGRAY
        r.font.name = "Arial"
        add_paragraph(doc, vendedor, font_size=10, color=DARK, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        add_paragraph(doc, "Consultor Comercial", font_size=8.5, color=LGRAY,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # Aceite do cliente
    add_paragraph(doc, "", space_after=20)
    add_paragraph(doc, "De acordo:", font_size=10, color=TEXT_COLOR,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("_" * 50)
    r.font.size = Pt(10)
    r.font.color.rgb = LGRAY
    r.font.name = "Arial"

    add_paragraph(doc, f"{dados['nome']}", font_size=10, color=DARK, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "Cliente", font_size=8.5, color=LGRAY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # ===== FOOTER =====
    add_horizontal_line(doc, "E8E8E8", 1)
    add_paragraph(
        doc,
        "Azevedo Contabilidade | CNPJ 35.304.872/0001-28 | www.azevedocontabilidade.com.br",
        font_size=7, color=LGRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4
    )

    # ===== SAVE TO BUFFER =====
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ===== CLI MODE (for testing) =====
if __name__ == "__main__":
    dados = {
        "tratamento": "Ao Sr.",
        "nome": "João da Silva",
        "telefone": "(84) 99999-0000",
        "email": "joao@email.com",
        "vendedor": "Allan",
        "introducao": "gestão contábil e fiscal da empresa SILVA LTDA",
        "servicos": [
            {"descricao": "Gestão contábil mensal", "valor": 800, "periodicidade": "Mensal"},
            {"descricao": "Abertura de empresa", "valor": 1200, "periodicidade": "Única Vez"},
        ],
        "desconto_pct": 0.10,
        "pix_cnpj": "33.540.066/0001-23",
        "pix_titular": "ALLAN SAYURE DE AZEVEDO BARBOSA",
        "observacao": "O valor refere-se exclusivamente aos serviços descritos acima.",
        "incluir_doc": True,
        "texto_doc": "Cópia da identidade do responsável de cada empresa para fazermos a procuração necessária.",
        "logo_path": os.path.join(os.path.dirname(__file__), "logo.png")
    }

    docx_bytes = gerar_docx(dados)
    output_path = "Proposta_Teste.docx"
    with open(output_path, "wb") as f:
        f.write(docx_bytes)
    print(f"Proposta gerada: {output_path} ({len(docx_bytes):,} bytes)")
